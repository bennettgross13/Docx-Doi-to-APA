import docx
import re
from subprocess import PIPE, Popen

"""Funtions"""


def MakeShortReference(str):
	year = ""
	count = 1
	for i in range(len(str)):
		if str[i].isdigit() and count <= 4:
			year += str[i]
			count += 1
	author = str.split(",")[0][:]
	ref = f"{author} et al., {year}"
	return ref


def cmdline(command):
	process = Popen(
		args=command,
		stdout=PIPE,
		shell=True
	)
	return process.communicate()[0]

def RemoveUrl(citation):
	if citation is not None:
		https = re.compile("https")
		https_loc = https.search(citation)
		https_start = https_loc.span(0)[0]
		citation = citation[:https_start]

	else:
		citation = "No citation found"
	return citation




def DoiToApa(doi):
	if doi is not None:
		doi_url = f'https://doi.org/{doi}'

		command = """curl -LH "Accept: text/x-bibliography; style=apa" """ + doi_url
		print(doi_url)

		citation = cmdline(command).decode()

	elif doi is None:
		citation = "Doi was not found"
	return citation


def ParagraphReplace(paragraph_num, doc):
	paragraph = doc.paragraphs[paragraph_num]

	for x in range(len(paragraph.runs)):
		if len(paragraph.text) < 3:
			continue
		run = paragraph.runs[x]


		klammer_offen, klammer_zu = re.compile("{"), re.compile("}")
		klammer_location_offen, klammer_location_zu = klammer_offen.search(run.text), klammer_zu.search(run.text)

		if klammer_location_zu is None or klammer_location_offen is None:

			continue
		else:

			start = klammer_location_offen.span(0)[0]
			end = klammer_location_zu.span(0)[0]

			doi_in_docx = run.text[start + 1: end].strip()

			citation = DoiToApa(doi_in_docx)

			short_reference = MakeShortReference(citation)
			print(short_reference)
			run.text = run.text[:start] + "(" + short_reference + ")" + run.text[end + 1:]
			print(citation)

			return RemoveUrl(citation)


def ParagraphReplaceAuto(path):
	citationlist = []
	doc = docx.Document(path)
	amount_paragraphs = len(doc.paragraphs)

	print(f"Scanning {amount_paragraphs} paragraphs")
	print("loading...")

	for i in range(amount_paragraphs):
		if len(doc.paragraphs[i].text) < 1:
			continue
		else:
			cite = ParagraphReplace(i, doc)

			if cite is not None:
				citationlist.append(cite)

	print(citationlist)

	for cite in citationlist:
		journal = cite.split(").")[1].split(". ")[1].split(", ")[0]
		index_citation = cite.find(journal)
		end_citation = cite[index_citation + len(journal):]
		start_citation = cite[:index_citation]

		new_paragraph = doc.add_paragraph()
		start_add = new_paragraph.add_run(start_citation)
		journal_add = new_paragraph.add_run(journal)
		end_add = new_paragraph.add_run(end_citation)
		journal_add.font.italic = True
		doc.add_paragraph()

	doc.save("ReplacedDocx.docx")
	print("Document saved as: ReplacedDocx.docx")


ParagraphReplaceAuto("Multiple4.docx")
